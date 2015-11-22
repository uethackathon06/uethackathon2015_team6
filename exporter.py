#!/usr/bin/env python
# -*- coding: utf-8 -*-

import docx , json

document = docx.Document()

fi = open("../input.txt" , "r")

s = fi.read()

test = json.JSONDecoder().decode(s)

document.add_heading(test["header"] , 0)
test = test["questions"]

for i in range(len(test)):
	ss = u"Câu " + str(i+1) + "." + test[i]["text"]
	if (test[i]["text"][len(test[i]["text"])-1] != '\n'):
		ss += '\n'
	for j in range(len(test[i]["choices"])):
		ss += chr(ord('A') + j) + ". " + test[i]["choices"][j]["text"]
		if (test[i]["choices"][j]["text"][len(test[i]["choices"][j]["text"])-1] != '\n'):
			ss += '\n'
	document.add_paragraph(ss)
	
document.add_page_break()

ss = u"Đáp án:\n"
for i in range(len(test)):
	ss += str(i+1) + ". "
	for j in range(len(test[i]["choices"])):
		if (test[i]["choices"][j]["correct"]):
			ss += chr(ord("A") + j) + " "
	ss += '\n'

document.add_paragraph(ss)

document.save('../test.docx')